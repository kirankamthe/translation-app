from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import os
import tempfile
import time 
import pandas as pd
from docx import Document
from werkzeug.utils import secure_filename
from transformers import MarianMTModel, MarianTokenizer
import sqlite3

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = tempfile.gettempdir()
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 15 * 1024 * 1024  # 15MB limit

# Initialize SQLite database
def init_db():
    print("Initializing translation memory database...")
    conn = sqlite3.connect('translation_memory.db')
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS translations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        source_text TEXT NOT NULL,
        translated_text TEXT NOT NULL,
        source_lang TEXT NOT NULL,
        target_lang TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    conn.commit()
    conn.close()
    print("Database initialized.")

init_db()

# Check memory
def check_translation_memory(text, source_lang, target_lang):
    conn = sqlite3.connect('translation_memory.db')
    cursor = conn.cursor()
    cursor.execute(
        "SELECT translated_text FROM translations WHERE source_text = ? AND source_lang = ? AND target_lang = ?",
        (text, source_lang, target_lang)
    )
    result = cursor.fetchone()
    conn.close()
    return result[0] if result else None

# Save to memory
def save_to_translation_memory(source_text, translated_text, source_lang, target_lang):
    conn = sqlite3.connect('translation_memory.db')
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO translations (source_text, translated_text, source_lang, target_lang) VALUES (?, ?, ?, ?)",
        (source_text, translated_text, source_lang, target_lang)
    )
    conn.commit()
    conn.close()

# Model cache
model_cache = {}

# Preload supported model at app startup
def preload_models():
    supported_models = [("en", "de")]
    for src, tgt in supported_models:
        model_name = f"Helsinki-NLP/opus-mt-{src}-{tgt}"
        print(f"Preloading model: {model_name}")
        tokenizer = MarianTokenizer.from_pretrained(model_name)
        model = MarianMTModel.from_pretrained(model_name)
        model_cache[model_name] = (tokenizer, model)
        print("✅ Model preloading complete. Ready to translate.")

preload_models()

def get_model(source_lang, target_lang):
    model_name = f"Helsinki-NLP/opus-mt-{source_lang}-{target_lang}"
    if model_name not in model_cache:
        print(f"Loading model: {model_name}")
        tokenizer = MarianTokenizer.from_pretrained(model_name)
        model = MarianMTModel.from_pretrained(model_name, from_tf=True)
        model_cache[model_name] = (tokenizer, model)
    return model_cache[model_name]

# Batch translation
def translate_batch(texts, source_lang, target_lang):
    tokenizer, model = get_model(source_lang, target_lang)
    print("Tokenizing input texts...")

    results = []
    to_translate = []
    cached_map = {}

    for text in texts:
        cached = check_translation_memory(text, source_lang, target_lang)
        if cached:
            cached_map[text] = cached
        else:
            to_translate.append(text)

    if to_translate:
        inputs = tokenizer(to_translate, return_tensors="pt", padding=True, truncation=True)
        print("Generating translations...")
        outputs = model.generate(
            **inputs,
            num_beams=1,          # Greedy search (1 beam, faster)
            no_repeat_ngram_size=2,  # Prevents repeating n-grams
        )
        print("Decoding translations...")
        translations = tokenizer.batch_decode(outputs, skip_special_tokens=True)
        print("Translations generated.")
        for src, tgt in zip(to_translate, translations):
            save_to_translation_memory(src, tgt, source_lang, target_lang)
            cached_map[src] = tgt

    for text in texts:
        results.append(cached_map[text])
    return results

# DOCX Translation
def translate_docx(file_path, source_lang, target_lang):
    print(f"Translating DOCX file: {file_path}")
    doc = Document(file_path)
    texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    translations = translate_batch(texts, source_lang, target_lang)

    idx = 0
    for para in doc.paragraphs:
        if para.text.strip():
            para.text = translations[idx]
            idx += 1

    translated_path = file_path.replace('.docx', '_translated.docx')
    doc.save(translated_path)
    print(f"Saved translated DOCX to: {translated_path}")
    return translated_path

# XLSX Translation
def translate_xlsx(file_path, source_lang, target_lang):
    print(f"Translating XLSX file: {file_path}")
    df = pd.read_excel(file_path)
    for col in df.columns:
        print(f"Translating column: {col}")
        texts = df[col].apply(lambda x: x.strip() if isinstance(x, str) else None)
        to_translate = texts.dropna().tolist()
        translations = translate_batch(to_translate, source_lang, target_lang)
        mapping = dict(zip(to_translate, translations))
        df[col] = df[col].apply(lambda x: mapping.get(x.strip(), x) if isinstance(x, str) else x)

    translated_path = file_path.replace('.xlsx', '_translated.xlsx')
    df.to_excel(translated_path, index=False)
    print(f"Saved translated XLSX to: {translated_path}")
    return translated_path

@app.route('/api/translate', methods=['POST'])
def translate_file():
    start_time = time.time()  # Start timing
    print("Request received at /api/translate")
    if 'file' not in request.files:
        print("No file uploaded.")
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if not file.filename:
        print("Empty filename.")
        return jsonify({'error': 'Empty filename'}), 400

    filename = secure_filename(file.filename)
    source_lang = request.form.get('sourceLanguage', 'en')
    target_lang = request.form.get('targetLanguage', 'fr')

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    print(f"File saved to: {file_path} | Source: {source_lang} | Target: {target_lang}")

    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.docx':
            translated_path = translate_docx(file_path, source_lang, target_lang)
        elif ext == '.xlsx':
            translated_path = translate_xlsx(file_path, source_lang, target_lang)
        else:
            print("Unsupported file format.")
            return jsonify({'error': 'Unsupported file format. Only .docx and .xlsx are supported.'}), 400

        elapsed_time = time.time() - start_time
        print(f"Translation completed in {elapsed_time:.2f} seconds.")
        return send_file(translated_path, as_attachment=True)

    except Exception as e:
        print(f"Translation failed: {e}")
        return jsonify({'error': str(e)}), 500

    finally:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"Cleaned up uploaded file: {file_path}")

if __name__ == '__main__':
    print("Starting Flask server on http://localhost:5000")
    app.run(debug=True, port=5000)
